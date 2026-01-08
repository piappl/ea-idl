"""Format-specific exporters and parsers for notes (YAML and DOCX)."""

import json
from typing import List

import yaml
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from eaidl.notes_model import NoteMetadata, NotesExport, NoteType


class YamlFormatter:
    """Exports/imports notes to/from YAML format."""

    @staticmethod
    def export(notes_export: NotesExport, output_path: str):
        """Export NotesExport to YAML file."""
        # Convert to dict structure for YAML
        export_data = {
            "metadata": {
                "export_timestamp": notes_export.metadata.export_timestamp.isoformat(),
                "root_packages": notes_export.metadata.root_packages,
                "database_url": notes_export.metadata.database_url,
                "note_count": notes_export.metadata.note_count,
            },
            "instructions": YamlFormatter._get_instructions(),
            "notes": [
                {
                    "type": note.note_type.value,
                    "object_id": note.object_id,
                    "note_id": note.note_id,
                    "checksum": note.checksum,
                    "path": note.path,
                    "namespace": note.namespace,
                    "object_name": note.object_name,
                    "object_guid": note.object_guid,
                    "content": note.content_md,
                }
                for note in sorted(notes_export.notes, key=lambda n: n.path)
            ],
        }

        # Write to YAML file with custom formatting
        with open(output_path, "w", encoding="utf-8") as f:
            # Write header comment
            f.write("# EA-IDL Notes Export (YAML format)\n")
            f.write("# \n")
            f.write("# EDITING INSTRUCTIONS:\n")
            f.write("# - You can edit the 'content' field of each note\n")
            f.write("# - Use markdown formatting: **bold**, *italic*, bullet lists\n")
            f.write(
                "# - DO NOT modify: type, object_id, note_id, checksum, path, namespace, object_name, object_guid\n"
            )
            f.write("# - Empty content is OK (will clear the note in EA)\n")
            f.write("# - During import, only notes unchanged in EA since export will be updated\n")
            f.write("# \n")
            f.write("# See 'instructions' section below for more details.\n")
            f.write("#\n\n")

            # Use custom YAML dumper for better formatting
            yaml.dump(
                export_data,
                f,
                default_flow_style=False,
                allow_unicode=True,
                sort_keys=False,
                width=120,
                indent=2,
            )

    @staticmethod
    def parse(yaml_path: str) -> List[NoteMetadata]:
        """Parse YAML file and return list of NoteMetadata."""
        with open(yaml_path, "r", encoding="utf-8") as f:
            data = yaml.safe_load(f)

        notes = []
        for note_dict in data.get("notes", []):
            note = NoteMetadata(
                note_type=NoteType(note_dict.get("type", "")),
                object_id=note_dict.get("object_id", 0),
                note_id=note_dict.get("note_id"),
                namespace=note_dict.get("namespace", []),
                object_name=note_dict.get("object_name", ""),
                content_md=note_dict.get("content", ""),
                content_html="",  # Will be generated from markdown during import
                checksum=note_dict.get("checksum", ""),
                path=note_dict.get("path", ""),
                object_guid=note_dict.get("object_guid"),
            )
            notes.append(note)

        return notes

    @staticmethod
    def _get_instructions() -> str:
        """Get editing instructions for reviewers."""
        return """
WHAT YOU CAN EDIT:
- The 'content' field of each note (markdown text)
- You can use markdown formatting: **bold**, *italic*, bullet lists

WHAT YOU MUST NOT EDIT:
- type, object_id, note_id, checksum (needed for validation)
- path, namespace, object_name (used for identification)
- object_guid (unique identifier for attributes)
- The metadata section
- The structure of the YAML file (don't add/remove notes)

FORMATTING GUIDE:
- **bold text** = Bold
- *italic text* = Italic
- * bullet item = Bullet list
- 1. numbered item = Numbered list

PARALLEL REVIEW:
- Multiple people can edit different sections of this file
- During import, only notes unchanged in EA since export will be updated
- Changed sections will be skipped (you'll get a report)

IMPORTANT:
- Do NOT delete note entries from the YAML
- Empty content is OK (will clear the note in EA)
- Keep all metadata fields intact for each note

Save this file when done and use it for import.
        """.strip()


class DocxFormatter:
    """Exports/imports notes to/from DOCX format."""

    @staticmethod
    def export(notes_export: NotesExport, output_path: str):
        """Export NotesExport to DOCX file."""
        doc = Document()

        # Add header
        DocxFormatter._add_header(doc, notes_export)

        # Add metadata table
        DocxFormatter._add_metadata_table(doc, notes_export)

        # Add instructions
        DocxFormatter._add_instructions(doc)

        # Add notes
        DocxFormatter._add_notes(doc, notes_export.notes)

        doc.save(output_path)

    @staticmethod
    def parse(docx_path: str) -> List[NoteMetadata]:
        """Parse DOCX file and return list of NoteMetadata."""
        doc = Document(docx_path)
        notes = []
        current_metadata = None
        current_content_lines = []

        # Iterate through document paragraphs and tables
        for element in doc.element.body:
            if element.tag.endswith("tbl"):  # Table
                table = DocxFormatter._get_table_for_element(doc, element)
                if table and DocxFormatter._is_metadata_table(table):
                    # Save previous note if exists
                    if current_metadata and current_content_lines:
                        current_metadata.content_md = "\n".join(current_content_lines).strip()
                        notes.append(current_metadata)
                        current_content_lines = []

                    # Parse new metadata
                    current_metadata = DocxFormatter._parse_metadata_table(table)

            elif element.tag.endswith("p"):  # Paragraph
                para = DocxFormatter._get_paragraph_for_element(doc, element)
                if para:
                    text = para.text.strip()

                    # Skip headings and special markers (but only if they have text)
                    if text and para.style.name.startswith("Heading"):
                        continue
                    if text.startswith("NOTE START"):
                        continue
                    if text.startswith("─"):  # Separator
                        continue

                    # Collect content (including blank lines for markdown formatting)
                    if current_metadata:
                        current_content_lines.append(text)

        # Don't forget last note
        if current_metadata and current_content_lines:
            current_metadata.content_md = "\n".join(current_content_lines).strip()
            notes.append(current_metadata)

        return notes

    @staticmethod
    def _add_header(doc: Document, notes_export: NotesExport):
        """Add document title and export info."""
        title = doc.add_heading("EA-IDL Notes Export", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        info = doc.add_paragraph()
        info.add_run(f"Exported: {notes_export.metadata.export_timestamp.isoformat()}\n").bold = True
        info.add_run(f"Root Packages: {', '.join(notes_export.metadata.root_packages)}\n")
        info.add_run(f"Total Notes: {notes_export.metadata.note_count}\n")

    @staticmethod
    def _add_metadata_table(doc: Document, notes_export: NotesExport):
        """Add hidden metadata table for round-trip validation."""
        doc.add_heading("Metadata (Do Not Edit)", level=1)

        metadata = {
            "export_timestamp": notes_export.metadata.export_timestamp.isoformat(),
            "root_packages": notes_export.metadata.root_packages,
            "database_url": notes_export.metadata.database_url,
            "note_count": notes_export.metadata.note_count,
        }

        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        row = table.rows[0]
        row.cells[0].text = "Export Metadata"
        row.cells[1].text = json.dumps(metadata, indent=2)

        # Make table cells smaller
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    @staticmethod
    def _add_instructions(doc: Document):
        """Add editing instructions for reviewers."""
        doc.add_page_break()
        doc.add_heading("Instructions for Reviewers", level=1)

        instructions = """
This document contains documentation notes from the EA model for review.

WHAT YOU CAN EDIT:
- Note content (text beneath each "NOTE START" marker)
- You can use markdown formatting: **bold**, *italic*, bullet lists

WHAT YOU MUST NOT EDIT:
- Section headings (package/class/attribute names)
- The metadata tables (small tables with Object ID, Note ID, Checksum, etc.)
- The structure of the document (adding/removing sections)

FORMATTING GUIDE:
- **bold text** = Bold
- *italic text* = Italic
- * bullet item = Bullet list
- 1. numbered item = Numbered list

PARALLEL REVIEW:
- Multiple people can edit different sections of this document
- During import, only notes unchanged in EA since export will be updated
- Changed sections will be skipped (you'll get a report)

IMPORTANT:
- Do NOT delete note sections
- Empty notes are OK (will clear the note in EA)
- Keep the metadata tables intact for each note

Save this file when done and send it back for import.
        """

        doc.add_paragraph(instructions.strip())

    @staticmethod
    def _add_notes(doc: Document, notes: List[NoteMetadata]):
        """Add all notes in hierarchical structure."""
        doc.add_page_break()
        doc.add_heading("Notes for Review", level=1)

        # Group notes by package
        current_package = None
        current_class = None

        for note in sorted(notes, key=lambda n: n.path):
            # Determine structure level
            if note.note_type in (NoteType.PACKAGE_MAIN, NoteType.PACKAGE_UNLINKED):
                if current_package != note.path.split("/")[0] if "/" in note.path else note.path:
                    current_package = note.path.split("/")[0] if "/" in note.path else note.path
                    current_class = None
                DocxFormatter._add_package_note(doc, note)
            elif note.note_type in (NoteType.CLASS_MAIN, NoteType.CLASS_LINKED):
                if current_class != note.path:
                    current_class = note.path
                DocxFormatter._add_class_note(doc, note)
            elif note.note_type in (NoteType.ATTRIBUTE_MAIN, NoteType.ATTRIBUTE_LINKED):
                DocxFormatter._add_attribute_note(doc, note)

    @staticmethod
    def _add_package_note(doc: Document, note: NoteMetadata):
        """Add a package note section."""
        if note.note_type == NoteType.PACKAGE_MAIN:
            heading_text = f"Package: {note.object_name}"
        else:
            heading_text = f"Package Note: {note.object_name} (unlinked #{note.note_id})"

        doc.add_heading(heading_text, level=2)
        DocxFormatter._add_note_metadata_table(doc, note)
        DocxFormatter._add_note_content(doc, note)

    @staticmethod
    def _add_class_note(doc: Document, note: NoteMetadata):
        """Add a class note section."""
        if note.note_type == NoteType.CLASS_MAIN:
            heading_text = f"Class: {note.object_name}"
        else:
            heading_text = f"Class Linked Note: {note.object_name} (#{note.note_id})"

        doc.add_heading(heading_text, level=3)
        DocxFormatter._add_note_metadata_table(doc, note)
        DocxFormatter._add_note_content(doc, note)

    @staticmethod
    def _add_attribute_note(doc: Document, note: NoteMetadata):
        """Add an attribute note section."""
        if note.note_type == NoteType.ATTRIBUTE_MAIN:
            heading_text = f"Attribute: {note.object_name}"
        else:
            heading_text = f"Attribute Linked Note: {note.object_name} (#{note.note_id})"

        doc.add_heading(heading_text, level=4)
        DocxFormatter._add_note_metadata_table(doc, note)
        DocxFormatter._add_note_content(doc, note)

    @staticmethod
    def _add_note_metadata_table(doc: Document, note: NoteMetadata):
        """Add metadata for a note (for round-trip validation)."""
        num_rows = 6 if note.object_guid else 5
        table = doc.add_table(rows=num_rows, cols=2)
        table.style = "Light Shading Accent 1"

        table.rows[0].cells[0].text = "Type"
        table.rows[0].cells[1].text = note.note_type.value

        table.rows[1].cells[0].text = "Object ID"
        table.rows[1].cells[1].text = str(note.object_id)

        table.rows[2].cells[0].text = "Note ID"
        table.rows[2].cells[1].text = str(note.note_id) if note.note_id else "N/A"

        table.rows[3].cells[0].text = "Checksum"
        table.rows[3].cells[1].text = note.checksum

        table.rows[4].cells[0].text = "Path"
        table.rows[4].cells[1].text = note.path

        if note.object_guid:
            table.rows[5].cells[0].text = "Object GUID"
            table.rows[5].cells[1].text = note.object_guid

        # Make table small
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    @staticmethod
    def _add_note_content(doc: Document, note: NoteMetadata):
        """Add the editable note content with markdown formatting."""
        # Add marker
        marker = doc.add_paragraph()
        marker.add_run("NOTE START (edit below):").bold = True

        # Add content paragraph
        content_para = doc.add_paragraph()
        content_para.add_run(note.content_md)

        # Add visual separator
        doc.add_paragraph("─" * 80)

    @staticmethod
    def _get_table_for_element(doc: Document, element):
        """Get table object from XML element."""
        for table in doc.tables:
            if table._element == element:
                return table
        return None

    @staticmethod
    def _get_paragraph_for_element(doc: Document, element):
        """Get paragraph object from XML element."""
        for para in doc.paragraphs:
            if para._element == element:
                return para
        return None

    @staticmethod
    def _is_metadata_table(table) -> bool:
        """Check if table is a metadata table."""
        if len(table.rows) < 5:
            return False
        # Check for "Type" in first row
        return "Type" in table.rows[0].cells[0].text or "Object ID" in table.rows[1].cells[0].text

    @staticmethod
    def _parse_metadata_table(table) -> NoteMetadata:
        """Parse metadata from table."""
        metadata = {}
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                metadata[key] = value

        # Parse note_id (handle "N/A")
        note_id = metadata.get("Note ID", "N/A")
        if note_id == "N/A":
            note_id = None
        else:
            note_id = int(note_id)

        # Parse object_guid (optional - only present for attributes)
        object_guid = metadata.get("Object GUID")
        if object_guid == "":
            object_guid = None

        # Extract object_name from path (last component)
        path = metadata.get("Path", "unknown")
        object_name = path.split("/")[-1] if "/" in path else path

        return NoteMetadata(
            note_type=NoteType(metadata.get("Type", "")),
            object_id=int(metadata.get("Object ID", 0)),
            note_id=note_id,
            namespace=[],  # Namespace is not stored in DOCX metadata
            object_name=object_name,
            content_md="",  # Will be filled from document
            content_html="",  # Will be generated from markdown
            checksum=metadata.get("Checksum", ""),
            path=path,
            object_guid=object_guid,
        )
