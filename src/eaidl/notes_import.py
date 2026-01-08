"""Import edited notes from DOCX back to EA database."""

from dataclasses import dataclass
from typing import List, Optional
from enum import Enum
import hashlib

from docx import Document
import markdown

from eaidl.config import Configuration
from eaidl.load import ModelParser, base
from eaidl.html_utils import strip_html, convert_to_ea_html
from eaidl.notes_export import NoteMetadata


class ImportStatus(Enum):
    """Status of a note import."""

    SUCCESS = "success"
    CHECKSUM_MISMATCH = "checksum_mismatch"
    NOT_FOUND = "not_found"
    CONTENT_UNCHANGED = "unchanged"
    ERROR = "error"


@dataclass
class NoteImportResult:
    """Result of importing a single note."""

    note_type: str
    object_id: int
    note_id: Optional[int]
    path: str
    status: ImportStatus
    message: str
    old_content: Optional[str] = None
    new_content: Optional[str] = None
    object_guid: Optional[str] = None


@dataclass
class ImportSummary:
    """Summary of entire import operation."""

    total_notes: int
    imported: int
    skipped_checksum: int
    skipped_unchanged: int
    not_found: int
    errors: int
    results: List[NoteImportResult]

    def print_report(self):
        """Print human-readable import report."""
        print("\n" + "=" * 80)
        print("NOTE IMPORT SUMMARY")
        print("=" * 80)
        print(f"Total notes in document: {self.total_notes}")
        print(f"Successfully imported:   {self.imported}")
        print(f"Skipped (checksum):      {self.skipped_checksum}")
        print(f"Skipped (unchanged):     {self.skipped_unchanged}")
        print(f"Not found in model:      {self.not_found}")
        print(f"Errors:                  {self.errors}")
        print("=" * 80)

        if self.skipped_checksum > 0:
            print(f"\n⚠  {self.skipped_checksum} notes were skipped due to checksum mismatch")
            print("   (EA model changed since export)")

        if self.errors > 0:
            print(f"\n❌ {self.errors} errors occurred during import")

        if self.imported > 0:
            print(f"\n✓ Successfully imported {self.imported} note updates")


class DocxImporter:
    """Imports notes from DOCX with per-note validation."""

    def __init__(self, docx_path: str, config: Configuration, parser: ModelParser):
        self.docx_path = docx_path
        self.config = config
        self.parser = parser
        self.doc = Document(docx_path)
        self.results: List[NoteImportResult] = []

    def parse_document(self) -> List[NoteMetadata]:
        """Parse DOCX and extract notes with metadata."""
        notes = []
        current_metadata = None
        current_content_lines = []

        # Iterate through document paragraphs and tables
        for element in self.doc.element.body:
            if element.tag.endswith("tbl"):  # Table
                table = self._get_table_for_element(element)
                if table and self._is_metadata_table(table):
                    # Save previous note if exists
                    if current_metadata and current_content_lines:
                        current_metadata.content_md = "\n".join(current_content_lines).strip()
                        notes.append(current_metadata)
                        current_content_lines = []

                    # Parse new metadata
                    current_metadata = self._parse_metadata_table(table)

            elif element.tag.endswith("p"):  # Paragraph
                para = self._get_paragraph_for_element(element)
                if para:
                    text = para.text.strip()

                    # Skip headings and special markers (but only if they have text)
                    if text and para.style.name.startswith("Heading"):
                        continue
                    if text.startswith("NOTE START"):
                        continue
                    if text.startswith("─"):  # Separator
                        continue

                    # Collect content (including blank lines for markdown formatting)
                    if current_metadata:
                        current_content_lines.append(text)

        # Don't forget last note
        if current_metadata and current_content_lines:
            current_metadata.content_md = "\n".join(current_content_lines).strip()
            notes.append(current_metadata)

        return notes

    def _get_table_for_element(self, element):
        """Get table object from XML element."""
        for table in self.doc.tables:
            if table._element == element:
                return table
        return None

    def _get_paragraph_for_element(self, element):
        """Get paragraph object from XML element."""
        for para in self.doc.paragraphs:
            if para._element == element:
                return para
        return None

    def _is_metadata_table(self, table) -> bool:
        """Check if table is a metadata table."""
        if len(table.rows) < 5:
            return False
        # Check for "Type" in first row
        return "Type" in table.rows[0].cells[0].text or "Object ID" in table.rows[1].cells[0].text

    def _parse_metadata_table(self, table) -> NoteMetadata:
        """Parse metadata from table."""
        metadata = {}
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip()
                value = row.cells[1].text.strip()
                metadata[key] = value

        # Parse note_id (handle "N/A")
        note_id = metadata.get("Note ID", "N/A")
        if note_id == "N/A":
            note_id = None
        else:
            note_id = int(note_id)

        # Parse object_guid (optional - only present for attributes)
        object_guid = metadata.get("Object GUID")
        if object_guid == "":
            object_guid = None

        return NoteMetadata(
            note_type=metadata.get("Type", ""),
            object_id=int(metadata.get("Object ID", 0)),
            note_id=note_id,
            namespace=[],  # Will be populated during validation
            object_name="",  # Will be populated during validation
            content_md="",  # Will be filled from document
            content_html="",  # Will be generated from markdown
            checksum=metadata.get("Checksum", ""),
            path=metadata.get("Path", ""),
            object_guid=object_guid,
        )

    def validate_and_import(
        self, notes: List[NoteMetadata], dry_run: bool = True, strict: bool = False
    ) -> ImportSummary:
        """Validate each note and prepare imports.

        Args:
            notes: Notes parsed from DOCX
            dry_run: If True, don't commit changes
            strict: If True, fail entire import on any checksum mismatch
        """
        for note in notes:
            result = self._validate_note(note)
            self.results.append(result)

            if strict and result.status == ImportStatus.CHECKSUM_MISMATCH:
                raise ValueError(
                    f"Strict mode: Checksum mismatch for {note.path}\n" f"EA model changed since export. Cannot import."
                )

        # Generate summary
        summary = self._generate_summary()

        # Perform imports if not dry run
        if not dry_run:
            self._commit_imports()

        return summary

    def _validate_note(self, note: NoteMetadata) -> NoteImportResult:
        """Validate a single note against current EA database.

        Returns ImportStatus indicating what to do with this note.
        """
        # Query current note content from database
        current_html = self._get_current_note_html(note)

        if current_html is None:
            return NoteImportResult(
                note_type=note.note_type,
                object_id=note.object_id,
                note_id=note.note_id,
                path=note.path,
                status=ImportStatus.NOT_FOUND,
                message=f"Object not found in database: {note.path}",
                object_guid=note.object_guid,
            )

        # Calculate current checksum
        current_checksum = hashlib.md5(current_html.encode("utf-8")).hexdigest()

        # Compare with exported checksum
        if current_checksum != note.checksum:
            return NoteImportResult(
                note_type=note.note_type,
                object_id=note.object_id,
                note_id=note.note_id,
                path=note.path,
                status=ImportStatus.CHECKSUM_MISMATCH,
                message="Checksum mismatch: EA model changed since export",
                old_content=strip_html(current_html),
                object_guid=note.object_guid,
            )

        # Checksum matches - check if content changed in DOCX
        current_md = strip_html(current_html)
        if current_md.strip() == note.content_md.strip():
            return NoteImportResult(
                note_type=note.note_type,
                object_id=note.object_id,
                note_id=note.note_id,
                path=note.path,
                status=ImportStatus.CONTENT_UNCHANGED,
                message="Content unchanged",
                object_guid=note.object_guid,
            )

        # All validation passed - this note can be imported
        return NoteImportResult(
            note_type=note.note_type,
            object_id=note.object_id,
            note_id=note.note_id,
            path=note.path,
            status=ImportStatus.SUCCESS,
            message="Ready to import",
            old_content=current_md,
            new_content=note.content_md,
            object_guid=note.object_guid,
        )

    def _get_current_note_html(self, note: NoteMetadata) -> Optional[str]:
        """Query current note content from EA database."""
        TObject = base.classes.t_object
        TAttribute = base.classes.t_attribute

        if note.note_type == "package_main":
            # Query package note from t_object (package as object)
            obj = self.parser.session.query(TObject).filter(TObject.attr_object_id == note.object_id).scalar()
            return obj.attr_note if obj else None

        elif note.note_type in ("package_unlinked", "class_linked", "attribute_linked"):
            # Query linked/unlinked note from t_object
            note_obj = self.parser.session.query(TObject).filter(TObject.attr_object_id == note.note_id).scalar()
            return note_obj.attr_note if note_obj else None

        elif note.note_type == "class_main":
            # Query class note
            obj = self.parser.session.query(TObject).filter(TObject.attr_object_id == note.object_id).scalar()
            return obj.attr_note if obj else None

        elif note.note_type == "attribute_main":
            # Query attribute note using GUID (attr_object_id is not unique - it's the parent class ID)
            attr = self.parser.session.query(TAttribute).filter(TAttribute.attr_ea_guid == note.object_guid).scalar()
            return attr.attr_notes if attr else None

        return None

    def _commit_imports(self):
        """Commit successful imports to database."""
        for result in self.results:
            if result.status != ImportStatus.SUCCESS:
                continue

            # Convert markdown back to HTML
            html_content = markdown.markdown(result.new_content, extensions=["extra", "sane_lists"])

            # Convert to EA-compatible HTML format (<b> instead of <strong>, etc.)
            html_content = convert_to_ea_html(html_content)

            # Update database based on note type
            self._update_note_in_database(result, html_content)

        # Commit all changes
        self.parser.session.commit()

    def _update_note_in_database(self, result: NoteImportResult, html_content: str):
        """Update a single note in the EA database."""
        TObject = base.classes.t_object
        TAttribute = base.classes.t_attribute

        if result.note_type == "package_main":
            obj = self.parser.session.query(TObject).filter(TObject.attr_object_id == result.object_id).first()
            if obj:
                obj.attr_note = html_content

        elif result.note_type in ("package_unlinked", "class_linked", "attribute_linked"):
            note_obj = self.parser.session.query(TObject).filter(TObject.attr_object_id == result.note_id).first()
            if note_obj:
                note_obj.attr_note = html_content

        elif result.note_type == "class_main":
            obj = self.parser.session.query(TObject).filter(TObject.attr_object_id == result.object_id).first()
            if obj:
                obj.attr_note = html_content

        elif result.note_type == "attribute_main":
            # Use GUID for attribute lookup (attr_object_id is not unique - it's the parent class ID)
            attr = self.parser.session.query(TAttribute).filter(TAttribute.attr_ea_guid == result.object_guid).first()
            if attr:
                attr.attr_notes = html_content

    def _generate_summary(self) -> ImportSummary:
        """Generate import summary from results."""
        status_counts = {
            ImportStatus.SUCCESS: 0,
            ImportStatus.CHECKSUM_MISMATCH: 0,
            ImportStatus.CONTENT_UNCHANGED: 0,
            ImportStatus.NOT_FOUND: 0,
            ImportStatus.ERROR: 0,
        }

        for result in self.results:
            status_counts[result.status] += 1

        return ImportSummary(
            total_notes=len(self.results),
            imported=status_counts[ImportStatus.SUCCESS],
            skipped_checksum=status_counts[ImportStatus.CHECKSUM_MISMATCH],
            skipped_unchanged=status_counts[ImportStatus.CONTENT_UNCHANGED],
            not_found=status_counts[ImportStatus.NOT_FOUND],
            errors=status_counts[ImportStatus.ERROR],
            results=self.results,
        )
