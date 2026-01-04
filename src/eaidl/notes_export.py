"""Export EA model notes to DOCX format for editing by non-EA users."""

from dataclasses import dataclass
from typing import List, Optional
from datetime import datetime
import hashlib

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import json

from eaidl.config import Configuration
from eaidl.model import ModelPackage, ModelClass, ModelAttribute
from eaidl.html_utils import strip_html


@dataclass
class NoteMetadata:
    """Metadata for a single note in the export."""

    note_type: str  # "package_main", "package_linked", "class_main", "class_linked", "attribute_main", "attribute_linked", "package_unlinked"
    object_id: int  # ID of package/class/attribute
    note_id: Optional[int]  # ID of Note object (None for main notes)
    namespace: List[str]  # Full namespace path
    object_name: str  # Name of package/class/attribute
    content_md: str  # Markdown content
    content_html: str  # Original HTML (for checksum)
    checksum: str  # MD5 of HTML content
    path: str  # Hierarchical path for DOCX structure
    object_guid: Optional[str] = None  # GUID for attributes (attr_object_id is not unique)


@dataclass
class NotesExport:
    """Container for all notes being exported."""

    notes: List[NoteMetadata]
    export_timestamp: str
    root_packages: List[str]
    database_url: str  # For reference


class NotesCollector:
    """Collects all notes from model into exportable format."""

    def __init__(self, config: Configuration, model: List[ModelPackage]):
        self.config = config
        self.model = model
        self.notes: List[NoteMetadata] = []

    def collect_all_notes(self) -> NotesExport:
        """Walk entire model tree and collect all notes."""
        for package in self.model:
            if package.name == "ext":  # Skip generated ext package
                continue
            self._collect_package_notes(package)

        return NotesExport(
            notes=self.notes,
            export_timestamp=datetime.now().isoformat(),
            root_packages=self.config.root_packages,
            database_url=self.config.database_url,
        )

    def _collect_package_notes(self, package: ModelPackage, parent_path: str = ""):
        """Recursively collect notes from package tree."""
        path = f"{parent_path}/{package.name}" if parent_path else package.name

        # Package main note
        if package.notes:
            self._add_note(
                note_type="package_main",
                object_id=package.object_id,
                namespace=package.namespace,
                object_name=package.name,
                content_html=package.notes,
                path=path,
            )

        # Package unlinked notes (free-floating notes in package)
        for idx, unlinked_note in enumerate(package.unlinked_notes):
            self._add_note(
                note_type="package_unlinked",
                object_id=package.package_id,
                note_id=unlinked_note.note_id,
                namespace=package.namespace,
                object_name=f"unlinked_{idx + 1}",
                content_html=unlinked_note.content_html,
                content_md=unlinked_note.content,
                checksum=unlinked_note.checksum,
                path=f"{path}/unlinked_{idx + 1}",
            )

        # Class notes
        for cls in package.classes:
            self._collect_class_notes(cls, path)

        # Recurse to child packages
        for child in package.packages:
            self._collect_package_notes(child, path)

    def _collect_class_notes(self, cls: ModelClass, package_path: str):
        """Collect notes from a class and its attributes."""
        path = f"{package_path}/{cls.name}"

        # Class main note
        if cls.notes:
            self._add_note(
                note_type="class_main",
                object_id=cls.object_id,
                namespace=cls.namespace,
                object_name=cls.name,
                content_html=cls.notes,
                path=path,
            )

        # Class linked notes
        for idx, linked_note in enumerate(cls.linked_notes):
            self._add_note(
                note_type="class_linked",
                object_id=cls.object_id,
                note_id=linked_note.note_id,
                namespace=cls.namespace,
                object_name=cls.name,
                content_html=linked_note.content_html,
                content_md=linked_note.content,
                checksum=linked_note.checksum,
                path=f"{path}/linked_{idx + 1}",
            )

        # Attribute notes
        for attr in cls.attributes:
            self._collect_attribute_notes(attr, path, cls.namespace)

    def _collect_attribute_notes(self, attr: ModelAttribute, class_path: str, parent_namespace: List[str]):
        """Collect notes from an attribute."""
        path = f"{class_path}/{attr.name}"

        # Attribute main note
        if attr.notes:
            self._add_note(
                note_type="attribute_main",
                object_id=attr.attribute_id,
                namespace=parent_namespace,
                object_name=attr.name,
                content_html=attr.notes,
                path=path,
                object_guid=attr.guid,
            )

        # Attribute linked notes
        for idx, linked_note in enumerate(attr.linked_notes):
            self._add_note(
                note_type="attribute_linked",
                object_id=attr.attribute_id,
                note_id=linked_note.note_id,
                namespace=parent_namespace,
                object_name=attr.name,
                content_html=linked_note.content_html,
                content_md=linked_note.content,
                checksum=linked_note.checksum,
                path=f"{path}/linked_{idx + 1}",
                object_guid=attr.guid,
            )

    def _add_note(
        self,
        note_type: str,
        object_id: int,
        namespace: List[str],
        object_name: str,
        content_html: str,
        path: str,
        note_id: Optional[int] = None,
        content_md: Optional[str] = None,
        checksum: Optional[str] = None,
        object_guid: Optional[str] = None,
    ):
        """Add a note to the collection with metadata."""
        # If content_md and checksum not provided, compute them
        if content_md is None:
            content_md = strip_html(content_html)
        if checksum is None:
            checksum = hashlib.md5(content_html.encode("utf-8")).hexdigest()

        self.notes.append(
            NoteMetadata(
                note_type=note_type,
                object_id=object_id,
                note_id=note_id,
                namespace=namespace,
                object_name=object_name,
                content_md=content_md,
                content_html=content_html,
                checksum=checksum,
                path=path,
                object_guid=object_guid,
            )
        )


class DocxExporter:
    """Exports notes to DOCX format with metadata."""

    def __init__(self, notes_export: NotesExport):
        self.export = notes_export
        self.doc = Document()

    def export_to_file(self, output_path: str):
        """Generate DOCX file."""
        self._add_header()
        self._add_metadata_table()
        self._add_instructions()
        self._add_notes()

        self.doc.save(output_path)

    def _add_header(self):
        """Add document title and export info."""
        title = self.doc.add_heading("EA-IDL Notes Export", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        info = self.doc.add_paragraph()
        info.add_run(f"Exported: {self.export.export_timestamp}\n").bold = True
        info.add_run(f"Root Packages: {', '.join(self.export.root_packages)}\n")
        info.add_run(f"Total Notes: {len(self.export.notes)}\n")

    def _add_metadata_table(self):
        """Add hidden metadata table for round-trip validation."""
        self.doc.add_heading("Metadata (Do Not Edit)", level=1)

        metadata = {
            "export_timestamp": self.export.export_timestamp,
            "root_packages": self.export.root_packages,
            "database_url": self.export.database_url,
            "note_count": len(self.export.notes),
        }

        table = self.doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        row = table.rows[0]
        row.cells[0].text = "Export Metadata"
        row.cells[1].text = json.dumps(metadata, indent=2)

        # Make table cells smaller
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)

    def _add_instructions(self):
        """Add editing instructions for reviewers."""
        self.doc.add_page_break()
        self.doc.add_heading("Instructions for Reviewers", level=1)

        instructions = """
This document contains documentation notes from the EA model for review.

WHAT YOU CAN EDIT:
- Note content (text beneath each "NOTE START" marker)
- You can use markdown formatting: **bold**, *italic*, bullet lists

WHAT YOU MUST NOT EDIT:
- Section headings (package/class/attribute names)
- The metadata tables (small tables with Object ID, Note ID, Checksum, etc.)
- The structure of the document (adding/removing sections)

FORMATTING GUIDE:
- **bold text** = Bold
- *italic text* = Italic
- * bullet item = Bullet list
- 1. numbered item = Numbered list

PARALLEL REVIEW:
- Multiple people can edit different sections of this document
- During import, only notes unchanged in EA since export will be updated
- Changed sections will be skipped (you'll get a report)

IMPORTANT:
- Do NOT delete note sections
- Empty notes are OK (will clear the note in EA)
- Keep the metadata tables intact for each note

Save this file when done and send it back for import.
        """

        self.doc.add_paragraph(instructions.strip())

    def _add_notes(self):
        """Add all notes in hierarchical structure."""
        self.doc.add_page_break()
        self.doc.add_heading("Notes for Review", level=1)

        # Group notes by package
        current_package = None
        current_class = None

        for note in sorted(self.export.notes, key=lambda n: n.path):
            # Determine structure level
            if note.note_type.startswith("package"):
                if current_package != note.path.split("/")[0] if "/" in note.path else note.path:
                    current_package = note.path.split("/")[0] if "/" in note.path else note.path
                    current_class = None
                self._add_package_note(note)
            elif note.note_type.startswith("class"):
                if current_class != note.path:
                    current_class = note.path
                self._add_class_note(note)
            elif note.note_type.startswith("attribute"):
                self._add_attribute_note(note)

    def _add_package_note(self, note: NoteMetadata):
        """Add a package note section."""
        # Determine if this is main or unlinked
        if note.note_type == "package_main":
            heading_text = f"Package: {note.object_name}"
        else:
            heading_text = f"Package Note: {note.object_name} (unlinked #{note.note_id})"

        self.doc.add_heading(heading_text, level=2)
        self._add_note_metadata_table(note)
        self._add_note_content(note)

    def _add_class_note(self, note: NoteMetadata):
        """Add a class note section."""
        if note.note_type == "class_main":
            heading_text = f"Class: {note.object_name}"
        else:
            heading_text = f"Class Linked Note: {note.object_name} (#{note.note_id})"

        self.doc.add_heading(heading_text, level=3)
        self._add_note_metadata_table(note)
        self._add_note_content(note)

    def _add_attribute_note(self, note: NoteMetadata):
        """Add an attribute note section."""
        if note.note_type == "attribute_main":
            heading_text = f"Attribute: {note.object_name}"
        else:
            heading_text = f"Attribute Linked Note: {note.object_name} (#{note.note_id})"

        self.doc.add_heading(heading_text, level=4)
        self._add_note_metadata_table(note)
        self._add_note_content(note)

    def _add_note_metadata_table(self, note: NoteMetadata):
        """Add metadata for a note (for round-trip validation)."""
        # Small table with metadata (6 rows if GUID present, 5 otherwise)
        num_rows = 6 if note.object_guid else 5
        table = self.doc.add_table(rows=num_rows, cols=2)
        table.style = "Light Shading Accent 1"

        table.rows[0].cells[0].text = "Type"
        table.rows[0].cells[1].text = note.note_type

        table.rows[1].cells[0].text = "Object ID"
        table.rows[1].cells[1].text = str(note.object_id)

        table.rows[2].cells[0].text = "Note ID"
        table.rows[2].cells[1].text = str(note.note_id) if note.note_id else "N/A"

        table.rows[3].cells[0].text = "Checksum"
        table.rows[3].cells[1].text = note.checksum

        table.rows[4].cells[0].text = "Path"
        table.rows[4].cells[1].text = note.path

        if note.object_guid:
            table.rows[5].cells[0].text = "Object GUID"
            table.rows[5].cells[1].text = note.object_guid

        # Make table small
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    def _add_note_content(self, note: NoteMetadata):
        """Add the editable note content with markdown formatting."""
        # Add marker
        marker = self.doc.add_paragraph()
        marker.add_run("NOTE START (edit below):").bold = True

        # Add content paragraph
        content_para = self.doc.add_paragraph()
        content_para.add_run(note.content_md)

        # Add visual separator
        self.doc.add_paragraph("─" * 80)
